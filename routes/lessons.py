from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, current_app, send_file
from flask_login import login_required, current_user
from models import Topic, LessonPlan, LessonActivity, db
import json
from pathlib import Path
from datetime import datetime
from flask_wtf import FlaskForm
from wtforms import StringField, SelectField, IntegerField
from wtforms.validators import DataRequired
import google.generativeai as genai
import os
from dotenv import load_dotenv
import pdfkit
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS

# Load environment variables
load_dotenv()

# Configure Gemini
GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY')
genai.configure(api_key=GOOGLE_API_KEY)

lessons = Blueprint('lessons', __name__)

@lessons.app_template_filter('from_json')
def from_json(value):
    return json.loads(value) if value else {}

class LessonPlanForm:
    def __init__(self, formdata=None):
        self.subject = formdata.get('subject') if formdata else None
        self.grade = formdata.get('grade') if formdata else None
        self.topic = formdata.get('topic') if formdata else None
        self.duration = formdata.get('duration') if formdata else None
        self.time = formdata.get('time') if formdata else None
        self.term = formdata.get('term') if formdata else None

    def validate(self):
        return all([
            self.subject,
            self.grade,
            self.topic,
            self.duration,
            self.time,
            self.term
        ])

@lessons.route('/lessons')
@login_required
def list_lessons():
    lesson_plans = LessonPlan.query.filter_by(user_id=current_user.id).all()
    form = LessonPlanForm()  # Create form instance
    return render_template('lessons/list.html', lesson_plans=lesson_plans, form=form)


@lessons.route('/lessons/<int:lesson_id>')
@login_required
def view_lesson(lesson_id):
    lesson = LessonPlan.query.get_or_404(lesson_id)
    if lesson.user_id != current_user.id:
        flash('Access denied')
        return redirect(url_for('lessons.list_lessons'))
    return render_template('lessons/view.html', lesson=lesson)

@lessons.route('/lessons/generate', methods=['POST'])
@login_required
def generate_lesson():
    form = LessonPlanForm(request.form)
    if form.validate():
        try:
            # Get form data
            subject = form.subject
            grade = form.grade
            topic = form.topic
            duration = form.duration
            time = form.time
            term = form.term

            # Initialize Gemini model directly
            model = genai.GenerativeModel('gemini-1.5-flash')  

            # Create prompt for lesson plan generation
            prompt = f"""
            Create a detailed lesson plan for:
            Subject: {subject}
            Grade Level: {grade}
            Topic: {topic}
            Duration: {duration} minutes

            Provide ONLY the JSON response without any markdown formatting or code blocks.
            The response must be a valid JSON object with this exact structure:
            {{
                "title": "Lesson title",
                "objectives": ["List of learning objectives"],
                "materials": ["Required materials"],
                "introduction": {{
                    "duration": "time in minutes",
                    "activities": ["List of introduction activities"]
                }},
                "main_content": [
                    {{
                        "duration": "time in minutes",
                        "topic": "subtopic name",
                        "teaching_activities": ["List of teaching activities"],
                        "student_activities": ["List of student activities"]
                    }}
                ],
                "conclusion": {{
                    "duration": "time in minutes",
                    "activities": ["List of concluding activities"]
                }},
                "assessment": ["List of assessment methods"],
                "homework": "Homework assignment"
            }}

            Do not include any explanatory text, markdown formatting, or code block markers.
            Return only the JSON object.
            """

            # Generate lesson plan content using Gemini
            try:
                response = model.generate_content(prompt)
                if not response or not response.text:
                    raise ValueError("No response from AI model")
                
                # Clean the response text to ensure it's valid JSON
                cleaned_text = response.text.strip()
                # Remove markdown code block if present
                if cleaned_text.startswith('```'):
                    cleaned_text = cleaned_text.split('\n', 1)[1]  # Remove first line
                    cleaned_text = cleaned_text.rsplit('\n', 1)[0]  # Remove last line with ```
                if cleaned_text.startswith('json'):
                    cleaned_text = cleaned_text.split('\n', 1)[1]  # Remove json language identifier
                
                lesson_content = json.loads(cleaned_text)
                
            except json.JSONDecodeError as e:
                current_app.logger.error(f"JSON parsing error: {str(e)}, Response: {response.text}")
                flash('Error parsing AI response. Please try again.', 'danger')
                return redirect(url_for('lessons.list_lessons'))
            except Exception as e:
                current_app.logger.error(f"AI Generation error: {str(e)}")
                flash('Error generating lesson content. Please try again.', 'danger')
                return redirect(url_for('lessons.list_lessons'))

            # Create lesson plan with additional fields
            lesson_plan = LessonPlan(
                title=lesson_content['title'],
                subject=subject,
                grade_level=grade,
                topic_name=topic,
                duration=duration,
                time=time,
                term=term,
                school=current_user.school,  # From user profile
                teacher=current_user.name,   # From user profile
                objectives=json.dumps(lesson_content['objectives']),
                materials=json.dumps(lesson_content['materials']),
                introduction=json.dumps(lesson_content['introduction']),
                main_content=json.dumps(lesson_content['main_content']),
                conclusion=json.dumps(lesson_content['conclusion']),
                assessment=json.dumps(lesson_content['assessment']),
                homework=lesson_content['homework'],
                user_id=current_user.id
            )

            # Add and commit to database
            try:
                db.session.add(lesson_plan)
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f"Database error: {str(e)}")
                flash('Error saving lesson plan. Please try again.', 'danger')
                return redirect(url_for('lessons.list_lessons'))

            flash('Lesson plan generated successfully!', 'success')
            return redirect(url_for('lessons.view_lesson', lesson_id=lesson_plan.id))

        except Exception as e:
            current_app.logger.error(f"Unexpected error: {str(e)}")
            flash('An unexpected error occurred. Please try again.', 'danger')
            return redirect(url_for('lessons.list_lessons'))

    # If form validation fails
    for field, errors in form.errors.items():
        for error in errors:
            flash(f'{field}: {error}', 'danger')
    return redirect(url_for('lessons.list_lessons'))

@lessons.route('/lessons/<int:lesson_id>/download/<format>')
@login_required
def download_lesson(lesson_id, format):
    lesson = LessonPlan.query.get_or_404(lesson_id)
    if lesson.user_id != current_user.id:
        flash('Access denied')
        return redirect(url_for('lessons.list_lessons'))

    try:
        if format == 'pdf':
            # Generate HTML content
            html_content = render_template(
                'lessons/pdf_template.html',
                lesson=lesson,
                current_user=current_user
            )
            
            # Create PDF using WeasyPrint
            html = HTML(string=html_content, base_url=request.url_root)
            pdf = html.write_pdf()
            
            # Create response
            stream = BytesIO(pdf)
            stream.seek(0)
            return send_file(
                stream,
                download_name=f'lesson_plan_{lesson.id}.pdf',
                mimetype='application/pdf',
                as_attachment=True
            )
            
        elif format == 'docx':
            # Create Word document
            doc = Document()
            
            # Add header with centered alignment
            header = doc.add_heading('LESSON PLAN', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section
            metadata_table = doc.add_table(rows=2, cols=2)
            metadata_table.style = 'Table Grid'
            
            # Left column
            left_cell = metadata_table.cell(0, 0)
            left_cell.text = f'School: {current_user.school}\nTeacher: {current_user.name}'
            left_cell = metadata_table.cell(1, 0)
            left_cell.text = f'Subject: {lesson.subject}\nTopic: {lesson.topic_name}'
            
            # Right column
            right_cell = metadata_table.cell(0, 1)
            right_cell.text = f'Grade: {lesson.grade_level}\nDuration: {lesson.duration} minutes'
            right_cell = metadata_table.cell(1, 1)
            right_cell.text = f'Time: {lesson.time}\nTerm: {lesson.term}'
            
            # Add sections
            doc.add_heading('Learning Objectives', level=1)
            objectives = json.loads(lesson.objectives)
            for obj in objectives:
                doc.add_paragraph(obj, style='List Bullet')
            
            doc.add_heading('Materials Needed', level=1)
            materials = json.loads(lesson.materials)
            for material in materials:
                doc.add_paragraph(material, style='List Bullet')
            
            # Introduction section
            intro = json.loads(lesson.introduction)
            doc.add_heading(f'Introduction ({intro["duration"]})', level=1)
            for activity in intro['activities']:
                doc.add_paragraph(activity, style='List Bullet')
            
            # Main Content
            doc.add_heading('Main Content', level=1)
            main_content = json.loads(lesson.main_content)
            for section in main_content:
                doc.add_heading(f'{section["topic"]} ({section["duration"]})', level=2)
                
                doc.add_paragraph('Teaching Activities:', style='Heading 3')
                for activity in section['teaching_activities']:
                    doc.add_paragraph(activity, style='List Bullet')
                
                doc.add_paragraph('Student Activities:', style='Heading 3')
                for activity in section['student_activities']:
                    doc.add_paragraph(activity, style='List Bullet')
            
            # Conclusion
            conclusion = json.loads(lesson.conclusion)
            doc.add_heading(f'Conclusion ({conclusion["duration"]})', level=1)
            for activity in conclusion['activities']:
                doc.add_paragraph(activity, style='List Bullet')
            
            # Assessment
            doc.add_heading('Assessment', level=1)
            assessment = json.loads(lesson.assessment)
            for method in assessment:
                doc.add_paragraph(method, style='List Bullet')
            
            # Homework
            doc.add_heading('Homework', level=1)
            doc.add_paragraph(lesson.homework)
            
            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            
            return send_file(
                docx_file,
                download_name=f'lesson_plan_{lesson.id}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True
            )
        
        flash('Invalid format requested')
        return redirect(url_for('lessons.view_lesson', lesson_id=lesson_id))
        
    except Exception as e:
        current_app.logger.error(f"Error downloading lesson plan: {str(e)}")
        flash('Error generating document. Please try again.', 'error')
        return redirect(url_for('lessons.view_lesson', lesson_id=lesson_id)) 