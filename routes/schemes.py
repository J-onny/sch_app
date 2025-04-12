from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, current_app, send_file
from flask_login import login_required, current_user
from models import SchemeOfWork, Topic
from extensions import db
from utils import generate_scheme_topics
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from weasyprint import HTML, CSS
from io import BytesIO
import tempfile
from docx.enum.text import WD_ALIGN_PARAGRAPH

schemes = Blueprint('schemes', __name__)

@schemes.route('/schemes')
@login_required
def list_schemes():
    schemes = SchemeOfWork.query.filter_by(user_id=current_user.id).all()
    return render_template('schemes/list.html', schemes=schemes)

@schemes.route('/schemes/new', methods=['GET', 'POST'])
@login_required
def create_scheme():
    if request.method == 'POST':
        title = request.form.get('title')
        subject = request.form.get('subject')
        grade_level = request.form.get('grade_level')
        term = request.form.get('term')
        
        scheme = SchemeOfWork(
            title=title,
            subject=subject,
            grade_level=grade_level,
            term=term,
            user_id=current_user.id
        )
        
        db.session.add(scheme)
        db.session.commit()
        
        # Generate topics using AI if requested
        if request.form.get('generate_topics'):
            ai_topics = generate_scheme_topics(subject, grade_level, term)
            if ai_topics:
                try:
                    topics_data = json.loads(ai_topics)
                    for i, topic_data in enumerate(topics_data['topics']):
                        topic = Topic(
                            title=topic_data['title'],
                            learning_objectives=topic_data['learning_objectives'],
                            duration=topic_data['duration'],
                            teaching_methods=topic_data['teaching_methods'],
                            assessment_methods=topic_data['assessment_methods'],
                            order=i+1,
                            scheme_id=scheme.id
                        )
                        db.session.add(topic)
                    db.session.commit()
                except Exception as e:
                    flash('Error processing AI-generated topics')
                    
        return redirect(url_for('schemes.view_scheme', scheme_id=scheme.id))
    
    return render_template('schemes/create.html')

@schemes.route('/schemes/<int:scheme_id>')
@login_required
def view_scheme(scheme_id):
    scheme = SchemeOfWork.query.get_or_404(scheme_id)
    if scheme.user_id != current_user.id:
        flash('Access denied')
        return redirect(url_for('schemes.list_schemes'))
    return render_template('schemes/view.html', scheme=scheme)

@schemes.route('/api/subjects')
@login_required
def get_subjects():
    syllabi_dir = Path(current_app.root_path) / 'static' / 'syllabi'
    # Get folder names instead of JSON files
    subjects = [d.name for d in syllabi_dir.iterdir() if d.is_dir()]
    return jsonify(subjects)

@schemes.route('/api/grades/<subject>')
@login_required
def get_grades(subject):
    syllabi_dir = Path(current_app.root_path) / 'static' / 'syllabi' / subject
    # Get all senior*.json files
    grades = [f.stem for f in syllabi_dir.glob('senior*.json')]
    return jsonify(grades)

@schemes.route('/api/syllabus/<subject>/<grade>')
@login_required
def get_syllabus(subject, grade):
    try:
        file_path = Path(current_app.root_path) / 'static' / 'syllabi' / subject / f'{grade}.json'
        with open(file_path) as f:
            syllabus = json.load(f)
            # Extract unique terms from the syllabus
            terms = sorted(set(item['term'] for item in syllabus))
            return jsonify({'terms': terms, 'content': syllabus})
    except FileNotFoundError:
        return jsonify({'error': 'Syllabus not found'}), 404

@schemes.route('/api/generate-scheme', methods=['POST'])
@login_required
def generate_scheme():
    data = request.get_json()
    
    try:
        # Generate scheme using AI
        scheme_data = generate_scheme_topics(
            subject=data['subject'],
            grade=data['grade'],
            term=data['term'],
            start_week=int(data['start_week']),
            end_week=int(data['end_week'])
        )
        
        if not scheme_data:
            raise ValueError("Failed to generate scheme topics")
        
        # Create scheme in database
        scheme = SchemeOfWork(
            title=f"{data['subject']} - {data['grade']} - Term {data['term']}",
            subject=data['subject'],
            grade_level=data['grade'],
            term=data['term'],
            duration=f"Week {data['start_week']} to Week {data['end_week']}",
            user_id=current_user.id
        )
        db.session.add(scheme)
        db.session.commit()
        
        # Add topics to database
        for i, topic_data in enumerate(scheme_data):
            topic = Topic(
                week=str(topic_data.get('week', '')),
                periods=int(topic_data.get('periods', 0)),
                theme=str(topic_data.get('theme', '')),
                title=str(topic_data.get('topic', '')),
                sub_topic=str(topic_data.get('sub_topic', '')),
                competency=str(topic_data.get('competency', '')),
                learning_objectives=json.dumps(topic_data.get('learning_outcomes', [])),
                learning_activities=json.dumps(topic_data.get('activities', [])),
                teaching_methods=json.dumps(topic_data.get('methodology', [])),
                resources=json.dumps(topic_data.get('resources', [])),
                references=json.dumps(topic_data.get('references', [])),
                remarks=str(topic_data.get('remarks', '')),
                order=i+1,
                scheme_id=scheme.id
            )
            db.session.add(topic)
            
        db.session.commit()
        return jsonify({'success': True, 'scheme_id': scheme.id})
        
    except Exception as e:
        current_app.logger.error(f"Error generating scheme: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})

@schemes.route('/schemes/<int:scheme_id>/download/<format>')
@login_required
def download_scheme(scheme_id, format):
    scheme = SchemeOfWork.query.get_or_404(scheme_id)
    if scheme.user_id != current_user.id:
        flash('Access denied')
        return redirect(url_for('schemes.list_schemes'))
    
    try:
        if format == 'pdf':
            # Use the PDF-specific template
            html_content = render_template(
                'schemes/pdf_template.html',
                scheme=scheme,
                current_user=current_user
            )
            
            # Create PDF using WeasyPrint
            html = HTML(string=html_content, base_url=request.url_root)
            pdf = html.write_pdf()
            
            # Create response
            stream = BytesIO(pdf)
            stream.seek(0)
            return send_file(
                stream,
                download_name=f'scheme_{scheme.id}.pdf',
                mimetype='application/pdf',
                as_attachment=True
            )
            
        elif format == 'docx':
            # Create Word document
            doc = Document()
            
            # Add header with centered alignment
            
            
            school_header = doc.add_heading(f'SCHOOL: {current_user.school.upper()}', 1)
            school_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section
            metadata_table = doc.add_table(rows=2, cols=2)
            metadata_table.style = 'Table Grid'
            metadata_table.autofit = False
            
            # Left column
            left_cell = metadata_table.cell(0, 0)
            left_cell.text = f'NAME OF TEACHER: {current_user.name.upper()}'
            left_cell = metadata_table.cell(1, 0)
            left_cell.text = f'SUBJECT: {scheme.subject.upper()}'
            
            # Right column with right alignment
            right_cell = metadata_table.cell(0, 1)
            right_cell.text = f'CLASS: {scheme.grade_level}'
            right_cell = metadata_table.cell(1, 1)
            right_cell.text = f'TERM: {scheme.term} YEAR: {scheme.created_at.year}'
            
            # Set right alignment for right column
            for row in metadata_table.rows:
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()  # Add spacing
            
            # Add main content table
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Set column widths (adjust these values as needed)
            widths = [0.5, 0.5, 1.5, 1.2, 1.5, 1.5, 1.2, 1.2, 1.0, 1.0]
            for i, width in enumerate(widths):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)
            
            # Add headers
            headers = ['WEEK', 'PERIODS', 'THEME/TOPIC\nSUB-TOPIC', 'COMPETENCY', 
                      'LEARNING OUTCOMES', 'LEARNING ACTIVITIES', 'METHODOLOGY', 
                      'TEACHING/LEARNING RESOURCES', 'REFERENCES', 'REMARKS']
            
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Make headers bold
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
            
            # Add topic rows
            for topic in scheme.topics:
                row = table.add_row().cells
                row[0].text = str(topic.week)
                row[1].text = str(topic.periods)
                
                # Theme/Topic cell with formatting
                theme_cell = row[2]
                theme_cell.text = ''  # Clear default text
                p = theme_cell.paragraphs[0]
                p.add_run(f'{topic.theme}\n').bold = True
                p.add_run(f'{topic.title}\n')
                p.add_run(topic.sub_topic).italic = True
                
                row[3].text = topic.competency
                
                # Add lists with bullet points
                row[4].text = '\n'.join(f'• {obj}' for obj in json.loads(topic.learning_objectives))
                row[5].text = '\n'.join(f'• {act}' for act in json.loads(topic.learning_activities))
                row[6].text = '\n'.join(f'• {method}' for method in json.loads(topic.teaching_methods))
                row[7].text = '\n'.join(f'• {res}' for res in json.loads(topic.resources))
                row[8].text = '\n'.join(f'• {ref}' for ref in json.loads(topic.references))
                row[9].text = topic.remarks
            
            # Add signature section
            doc.add_paragraph()  # Add spacing
            signatures_table = doc.add_table(rows=3, cols=3)
            signatures_table.autofit = False
            
            # Left signature
            left_sig = signatures_table.cell(0, 0)
            left_sig.text = 'Signed: _________________________'
            signatures_table.cell(1, 0).text = 'Director of Studies'
            signatures_table.cell(2, 0).text = 'Date: _________________________'
            
            # Right signature (last column)
            right_sig = signatures_table.cell(0, 2)
            right_sig.text = 'Signed: _________________________'
            signatures_table.cell(1, 2).text = 'Head of Department'
            signatures_table.cell(2, 2).text = 'Date: _________________________'
            
            # Right align the right signature column
            for i in range(3):
                signatures_table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            
            return send_file(
                docx_file,
                download_name=f'scheme_{scheme.id}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True
            )
        
        flash('Invalid format requested')
        return redirect(url_for('schemes.view_scheme', scheme_id=scheme_id))
        
    except Exception as e:
        current_app.logger.error(f"Error downloading scheme: {str(e)}")
        flash('Error generating document. Please try again.', 'error')
        return redirect(url_for('schemes.view_scheme', scheme_id=scheme_id))

@schemes.route('/<int:scheme_id>/delete', methods=['DELETE'])
@login_required
def delete_scheme(scheme_id):
    scheme = SchemeOfWork.query.get_or_404(scheme_id)
    if scheme.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
        
    try:
        db.session.delete(scheme)
        db.session.commit()
        return jsonify({'status': 'success'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500 