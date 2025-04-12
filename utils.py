import google.generativeai as genai
import os
import json
from dotenv import load_dotenv
from flask import current_app
from bs4 import BeautifulSoup
from pathlib import Path
from weasyprint import HTML, CSS
from jinja2 import Template
import tempfile
import logging
from models import Question, QuestionOption

# Try to import PDF/DOC generation libraries
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    PDF_ENABLED = True
except ImportError:
    PDF_ENABLED = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOC_ENABLED = True
except ImportError:
    DOC_ENABLED = False

from io import BytesIO

# Load environment variables from .env file
load_dotenv()

# Configure Gemini AI
GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY')

# Add at the top with other constants
LESSON_PLAN_TEMPLATE = {
    "metadata": {
        "school": "",
        "subject": "",
        "teacher": "",
        "class": "",
        "term": "",
        "date": "",
        "time": "",
        "duration": "",
        "learners": {
            "boys": 0,
            "girls": 0
        }
    },
    "content": {
        "theme": "",
        "topic": "",
        "subtopic": "",
        "competency": "",
        "learning_outcomes": [],
        "generic_skills": [],
        "values": [],
        "cross_cutting_issues": [],
        "key_learning_outcomes": []
    },
    "lesson_development": {
        "introduction": {
            "duration": "5 minutes",
            "teacher_activities": [],
            "learner_activities": [],
            "competencies": [],
            "resources": []
        },
        "presentation": [
            {
                "duration": "",
                "phase": "",
                "teacher_activities": [],
                "learner_activities": [],
                "competencies": [],
                "resources": []
            }
        ],
        "conclusion": {
            "duration": "5 minutes",
            "teacher_activities": [],
            "learner_activities": [],
            "competencies": [],
            "resources": []
        }
    },
    "evaluation": "",
    "references": []
}

# Add this near the top with other constants
SUBJECT_SCENARIO_TEMPLATES = {
    'Physics': {
        'laboratory': {
            'context': "In a physics laboratory experiment, {setup}",
            'types': ['measurement', 'analysis', 'error-calculation']
        },
        'real_world': {
            'context': "In a real-world application, {situation}",
            'types': ['mechanics', 'electricity', 'waves']
        },
        'engineering': {
            'context': "An engineering problem requires {problem}",
            'types': ['design', 'optimization', 'troubleshooting']
        }
    },
    'Mathematics': {
        'financial': {
            'context': "In a financial scenario, {situation}",
            'types': ['interest', 'investment', 'depreciation']
        },
        'geometric': {
            'context': "Consider a geometric problem where {setup}",
            'types': ['construction', 'proof', 'optimization']
        },
        'statistical': {
            'context': "Given a dataset from {domain}",
            'types': ['analysis', 'probability', 'inference']
        }
    },
    'Chemistry': {
        'laboratory': {
            'context': "During a chemistry experiment, {setup}",
            'types': ['titration', 'reaction', 'analysis']
        },
        'industrial': {
            'context': "In an industrial process, {process}",
            'types': ['manufacturing', 'quality-control', 'optimization']
        },
        'environmental': {
            'context': "In an environmental study, {situation}",
            'types': ['pollution', 'remediation', 'monitoring']
        }
    }
}

def get_ai_model():
    """Get the AI model, initializing it if needed"""
    try:
        if not GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY environment variable is not set. Please set it in your .env file.")
        
        if not hasattr(get_ai_model, 'model'):
            genai.configure(api_key=GOOGLE_API_KEY)
            get_ai_model.model = genai.GenerativeModel('gemini-1.5-flash')
        return get_ai_model.model
    except Exception as e:
        current_app.logger.error(f"Error initializing AI model: {str(e)}")
        raise

SYLLABI_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'syllabi')
os.makedirs(SYLLABI_DIR, exist_ok=True)

# Add to your existing imports and configuration
SCHEMES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'schemes')
os.makedirs(SCHEMES_DIR, exist_ok=True)

# Helper function for parsing scheme content
def parse_scheme_content(content):
    try:
        scheme_data = []
        if '<table' in content.lower():
            soup = BeautifulSoup(content, 'html.parser')  # Parse HTML content
            rows = soup.find_all('tr')
            
            for row in rows[1:]:  # Skip header row
                cells = row.find_all('td')
                if len(cells) == 6:  # Ensure correct column count
                    scheme_data.append({
                        'week': cells[0].text.strip(),
                        'topic': cells[1].text.strip(),
                        'content': cells[2].text.strip(),
                        'period': cells[3].text.strip(),
                        'learning_outcomes': cells[4].text.strip(),
                        'assessment_methods': cells[5].text.strip()
                    })
        return scheme_data
    except Exception as e:
        current_app.logger.error(f"Error parsing scheme content: {str(e)}")
        return []

def get_ai_response(term_topics, start_week, end_week):
    """Get response from AI model for scheme generation"""
    try:
        model = get_ai_model()
        
        # Create the prompt for the AI
        prompt = f"""
        Create a scheme of work for weeks {start_week} to {end_week}.
        Use these topics as reference:
        {json.dumps(term_topics, indent=2)}
        
        Return ONLY a JSON array. Each item must be an object with these fields:
        {{
            "week": number,
            "periods": number,
            "theme": string,
            "title": string,
            "competency": string,
            "learning_outcomes": [string],
            "learning_activities": [string],
            "resources": [string],
            "assessment_methods": [string]
        }}
        
        Important: Return ONLY the JSON array with no additional text or explanation.
        """
        
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        logging.debug(f"Raw AI response: {response_text[:200]}...")
        
        # Try to parse as JSON
        try:
            # First try direct parsing
            data = json.loads(response_text)
            if isinstance(data, list):
                return json.dumps(data)  # Return as JSON string
            else:
                raise ValueError("AI response is not a list")
                
        except json.JSONDecodeError:
            # Try to extract JSON array
            import re
            json_match = re.search(r'\[(.*)\]', response_text, re.DOTALL)
            if json_match:
                array_text = f"[{json_match.group(1)}]"
                # Validate it's proper JSON
                data = json.loads(array_text)
                if isinstance(data, list):
                    return json.dumps(data)  # Return as JSON string
                
            raise ValueError("Could not extract valid JSON array from AI response")
            
    except Exception as e:
        logging.error(f"Error in get_ai_response: {str(e)}")
        logging.error(f"AI response was: {response_text if 'response_text' in locals() else 'No response'}")
        raise

def generate_scheme(subject, grade, term, start_week, end_week, syllabus_data):
    try:
        logging.debug(f"Generating scheme for {subject} {grade} Term {term}")
        logging.debug(f"Syllabus data type: {type(syllabus_data)}")
        
        if not syllabus_data:
            # Load syllabus data from file if not provided
            syllabi_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'syllabi')
            syllabus_path = os.path.join(syllabi_dir, subject, f"{grade}.json")
            logging.debug(f"Loading syllabus from: {syllabus_path}")
            
            try:
                with open(syllabus_path) as f:
                    syllabus_data = json.load(f)
            except Exception as e:
                logging.error(f"Error loading syllabus file: {e}")
                raise ValueError(f"Could not load syllabus for {subject} {grade}")
        
        if not syllabus_data:
            raise ValueError("No syllabus data available")
            
        logging.debug(f"First syllabus item: {syllabus_data[0] if syllabus_data else 'No data'}")
        
        # Add type conversion for term to ensure consistency
        term = str(term)
        
        # Filter syllabus data for the specified term
        term_topics = [topic for topic in syllabus_data if str(topic['term']) == term]
        logging.debug(f"Filtered topics count: {len(term_topics)}")
        logging.debug(f"First filtered topic: {term_topics[0] if term_topics else 'No topics found'}")
        
        if not term_topics:
            raise ValueError(f"No topics found for term {term}")
        
        # Send to AI model and get response
        logging.debug("Sending prompt to AI model")
        response = get_ai_response(term_topics, start_week, end_week)
        
        logging.debug(f"AI Response type: {type(response)}")
        logging.debug(f"AI Response preview: {response[:200] if isinstance(response, str) else 'Non-string response'}")
        
        # Parse the response
        if isinstance(response, str):
            scheme_data = json.loads(response)
        else:
            scheme_data = response
            
        logging.debug(f"Parsed scheme data type: {type(scheme_data)}")
        logging.debug(f"First scheme item: {scheme_data[0] if isinstance(scheme_data, list) else 'Not a list'}")
        
        # Verify we have a list
        if not isinstance(scheme_data, list):
            raise ValueError(f"Expected scheme data to be a list, got {type(scheme_data)}")
            
        # Validate each item in the scheme
        for item in scheme_data:
            if not isinstance(item, dict):
                raise ValueError(f"Expected dict items in scheme, got {type(item)}")
                
        return scheme_data
        
    except Exception as e:
        logging.error(f"Error in generate_scheme: {str(e)}")
        raise

def generate_resource_recommendations(topic_title, subject):
    """Generate resource recommendations using Gemini AI"""
    prompt = f"""
    As an educational resource expert, recommend learning resources for teaching "{topic_title}" in {subject}.
    Return the response in JSON format with this structure:
    {{
        "resources": [
            {{
                "title": "Resource title",
                "description": "Brief description",
                "type": "video/article/interactive/worksheet",
                "url": "URL if applicable",
                "source": "Platform or provider name"
            }}
        ]
    }}
    Include a mix of different resource types (videos, articles, worksheets, etc.)
    Focus on high-quality, educational resources from reputable sources.
    """
    
    try:
        response = model.generate_content(prompt)
        return json.loads(response.text)
    except Exception as e:
        print(f"Error generating resource recommendations: {e}")
        return {"resources": []}

def generate_assessment(subject, grade_level, assessment_type, question_types):
    """Generate an assessment using AI."""
    try:
        # Define grade-specific requirements
        grade_requirements = {
            'senior1': """
                - Use basic introductory concepts
                - Focus on fundamental principles
                - Include simple calculations and definitions
                - Use straightforward language
                - Avoid complex scenarios
            """,
            'senior2': """
                - Build on senior1 concepts
                - Include moderate calculations
                - Introduce some analytical thinking
                - Use clear, age-appropriate scenarios
                - Connect topics to daily life
            """,
            'senior3': """
                - Include more complex analysis
                - Require deeper understanding
                - Include practical applications
                - Use more challenging calculations
                - Encourage critical thinking
            """,
            'senior4': """
                - Include advanced concepts
                - Require synthesis of multiple topics
                - Include complex problem-solving
                - Use real-world applications
                - Require detailed analysis
            """
        }.get(grade_level.lower(), """
            - Adjust difficulty to grade level
            - Use appropriate complexity
            - Include relevant examples
        """)

        # Create subject-specific instructions
        subject_instructions = {
            'Physics': f"""
            For {grade_level}:
            - Include age-appropriate physics applications
            - Use formulas appropriate for this level
            - Include diagrams where needed
            - Focus on concepts covered in {grade_level} syllabus
            """,
            'Chemistry': f"""
            For {grade_level}:
            - Include appropriate laboratory concepts
            - Use chemical equations suitable for this level
            - Focus on topics covered in {grade_level}
            - Include relevant safety considerations
            """,
            'Mathematics': f"""
            For {grade_level}:
            - Use mathematical concepts appropriate for this level
            - Include step-by-step solutions
            - Focus on topics in {grade_level} syllabus
            - Use clear mathematical notation
            """
        }.get(subject, f"""
            For {grade_level}:
            - Use grade-appropriate content
            - Include relevant examples
            - Match syllabus requirements
        """)
        
        # Create the base prompt
        prompt = f"""
        Create a {assessment_type} for {subject} at {grade_level} level.
        
        Grade Level Requirements:
        {grade_requirements}
        
        Subject Requirements:
        {subject_instructions}
        
        General Requirements:
        1. Questions must strictly match {grade_level} curriculum level
        2. Include a mix of: {', '.join(question_types)}
        3. Use vocabulary and concepts appropriate for {grade_level}
        4. Ensure scenarios are relevant to {grade_level} students
        5. Multiple choice questions should have 4 options with one correct answer
        6. Total marks should be appropriate for the duration
        
        Question Structure:
        - Start with grade-appropriate context
        - Use clear, level-appropriate language
        - For numerical questions:
          * Use calculations suitable for {grade_level}
          * Include appropriate units
          * Show solution steps matching grade level
        - For theory questions:
          * Link to syllabus topics
          * Use examples students can relate to
        
        Return the response in this JSON structure:
        {{
            "total_marks": 100,
            "questions": [
                {{
                    "question_text": "question here",
                    "context": "background information",
                    "question_type": "question type",
                    "marks": 10,
                    "correct_answer": "answer here",
                    "solution_steps": ["step 1", "step 2"],
                    "options": [
                        {{"text": "option 1", "is_correct": true}},
                        {{"text": "option 2", "is_correct": false}}
                    ],
                    "guide": {{
                        "common_mistakes": ["mistake 1", "mistake 2"],
                        "teaching_notes": ["note 1", "note 2"],
                        "rubric": [
                            {{
                                "name": "criterion",
                                "marks": 5,
                                "description": "details"
                            }}
                        ]
                    }}
                }}
            ]
        }}
        """
        
        model = get_ai_model()
        response = model.generate_content(prompt)
        
        # Log the raw response
        current_app.logger.info(f"Raw AI response: {response.text}")
        
        # Extract and validate JSON
        try:
            # First try direct JSON parsing
            ai_assessment = json.loads(response.text)
        except json.JSONDecodeError:
            # Try to extract JSON from markdown code block
            import re
            
            # Look for JSON content between ```json and ``` markers
            json_match = re.search(r'```json\s*(.*?)\s*```', response.text, re.DOTALL)
            if json_match:
                json_content = json_match.group(1)
                
                # Remove any trailing comments that might break JSON parsing
                comment_start = json_content.find('//')
                if comment_start != -1:
                    json_content = json_content[:comment_start]
                
                # Remove any trailing commas before closing brackets/braces
                json_content = re.sub(r',(\s*[}\]])', r'\1', json_content)
                
                try:
                    ai_assessment = json.loads(json_content)
                except json.JSONDecodeError as e:
                    current_app.logger.error(f"Failed to parse extracted JSON: {e}")
                    current_app.logger.error(f"Extracted content: {json_content}")
                    raise ValueError("Invalid JSON in AI response")
            else:
                # Try to find JSON between curly braces as last resort
                json_match = re.search(r'({[^}]*})', response.text)
                if json_match:
                    try:
                        ai_assessment = json.loads(json_match.group(1))
                    except json.JSONDecodeError:
                        raise ValueError("Could not parse JSON from response")
                else:
                    raise ValueError("No JSON found in response")

        # Validate the structure
        if not isinstance(ai_assessment, dict):
            raise ValueError(f"Expected dict, got {type(ai_assessment)}")
        if 'questions' not in ai_assessment:
            raise ValueError("Missing 'questions' key in response")
        if not isinstance(ai_assessment['questions'], list):
            raise ValueError("'questions' must be a list")

        # Clean up the data
        for question in ai_assessment['questions']:
            # Ensure required fields exist
            if 'question_text' not in question:
                raise ValueError("Question missing question_text")
            
            # Clean up question type
            question['question_type'] = question.get('question_type', 'short_answer').lower()
            
            # Ensure guide structure exists
            if 'guide' not in question:
                question['guide'] = {}
            
            # Ensure arrays exist even if empty
            question['guide']['common_mistakes'] = question.get('guide', {}).get('common_mistakes', [])
            question['guide']['teaching_notes'] = question.get('guide', {}).get('teaching_notes', [])
            question['guide']['rubric'] = question.get('guide', {}).get('rubric', [])

        return ai_assessment

    except Exception as e:
        current_app.logger.error(f"Error in generate_assessment: {str(e)}")
        current_app.logger.error(f"Response was: {response.text if 'response' in locals() else 'No response generated'}")
        raise

def parse_table_to_json(table):
    topics = []
    rows = table.find_all('tr')[1:]  # Skip header
    
    for row in rows:
        cells = row.find_all('td')
        if len(cells) >= 9:  # Ensure all columns present
            topic = {
                "week": cells[0].text.strip(),
                "periods": int(cells[1].text.strip().split()[0]),
                "theme": cells[2].text.strip().split('\n')[0],
                "topic": cells[2].text.strip().split('\n')[1],
                "competency": cells[3].text.strip(),
                "learning_outcomes": [li.strip() for li in cells[4].find_all('li')],
                "methodology": [li.strip() for li in cells[6].find_all('li')],
                "resources": [li.strip() for li in cells[7].find_all('li')],
                "references": [li.strip() for li in cells[7].find_all('li')],
                "remarks": cells[8].text.strip()
            }
            topics.append(topic)
    
    return topics

def get_local_guidelines(subject):
    """Load relevant local guidelines for lesson planning"""
    guidelines = {}
    base_path = Path(current_app.root_path) / 'static' / 'guidelines'
    
    # Load general guidelines
    with open(base_path / 'general' / 'competencies.json') as f:
        guidelines['competencies'] = json.load(f)
    
    # Load subject-specific guidelines
    subject_path = base_path / 'subjects' / subject.lower()
    if subject_path.exists():
        with open(subject_path / 'teaching_methods.json') as f:
            guidelines['teaching_methods'] = json.load(f)
    
    # Load national guidelines
    with open(base_path / 'national' / 'curriculum_objectives.json') as f:
        guidelines['national'] = json.load(f)
        
    return guidelines

def generate_guide_pdf(guide_data):
    """Generate a PDF version of the teacher's guide using WeasyPrint"""
    
    # Create HTML template
    html_template = """
    <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; }
                .header { text-align: center; margin-bottom: 20px; }
                .details { margin-bottom: 20px; }
                .details table { width: 100%; border-collapse: collapse; }
                .details td { padding: 8px; border: 1px solid #ddd; }
                .question { margin-bottom: 30px; }
                .question h2 { color: #2c3e50; }
                .answer { margin-left: 20px; }
                .mistakes, .notes { margin: 10px 0; padding-left: 20px; }
                .mistakes li, .notes li { margin: 5px 0; }
                @page { margin: 2.5cm; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{{ title }}</h1>
            </div>
            
            <div class="details">
                <table>
                    <tr>
                        <td><strong>Subject:</strong></td>
                        <td>{{ assessment.subject }}</td>
                        <td><strong>Grade Level:</strong></td>
                        <td>{{ assessment.grade_level }}</td>
                    </tr>
                    <tr>
                        <td><strong>Duration:</strong></td>
                        <td>{{ assessment.duration }}</td>
                        <td><strong>Total Marks:</strong></td>
                        <td>{{ assessment.total_marks }}</td>
                    </tr>
                </table>
            </div>
            
            {% for question in questions %}
            <div class="question">
                <h2>Question {{ loop.index }} ({{ question.marks }} marks)</h2>
                <p>{{ question.question_text }}</p>
                
                <div class="answer">
                    <h3>Model Answer:</h3>
                    <p>{{ question.correct_answer }}</p>
                </div>
                
                {% if question.common_mistakes %}
                <div class="mistakes">
                    <h3>Common Mistakes:</h3>
                    <ul>
                    {% for mistake in question.common_mistakes %}
                        <li>{{ mistake }}</li>
                    {% endfor %}
                    </ul>
                </div>
                {% endif %}
                
                {% if question.teaching_notes %}
                <div class="notes">
                    <h3>Teaching Notes:</h3>
                    <ul>
                    {% for note in question.teaching_notes %}
                        <li>{{ note }}</li>
                    {% endfor %}
                    </ul>
                </div>
                {% endif %}
            </div>
            {% endfor %}
        </body>
    </html>
    """
    
    # Render template with data
    template = Template(html_template)
    html_content = template.render(**guide_data)
    
    # Generate PDF
    html = HTML(string=html_content)
    return html.write_pdf()

def generate_guide_doc(guide_data):
    """Generate a Word document version of the teacher's guide"""
    if not DOC_ENABLED:
        raise ImportError("python-docx is not installed. Please install it using: pip install python-docx")
    
    doc = Document()
    
    # Title
    doc.add_heading(guide_data['title'], 0)

    # Assessment Details
    assessment = guide_data['assessment']
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    details = [
        ['Subject:', assessment.subject],
        ['Grade Level:', assessment.grade_level],
        ['Duration:', assessment.duration],
        ['Total Marks:', str(assessment.total_marks)]
    ]
    for detail in details:
        row = table.add_row().cells
        row[0].text = detail[0]
        row[1].text = detail[1]

    doc.add_paragraph()

    # Questions
    for i, question in enumerate(guide_data['questions'], 1):
        # Question Header
        doc.add_heading(f"Question {i} ({question['marks']} marks)", level=1)
        doc.add_paragraph(question['question_text'])
        
        # Model Answer
        doc.add_heading("Model Answer:", level=2)
        doc.add_paragraph(question['correct_answer'])
        
        # Common Mistakes
        if question.get('common_mistakes'):
            doc.add_heading("Common Mistakes:", level=2)
            for mistake in question['common_mistakes']:
                p = doc.add_paragraph()
                p.add_run("• ").bold = True
                p.add_run(mistake)
        
        # Teaching Notes
        if question.get('teaching_notes'):
            doc.add_heading("Teaching Notes:", level=2)
            for note in question['teaching_notes']:
                p = doc.add_paragraph()
                p.add_run("• ").bold = True
                p.add_run(note)
        
        doc.add_paragraph()  # Add spacing between questions

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_scheme_topics(subject, grade, term, start_week, end_week):
    """Generate scheme topics using AI"""
    try:
        # Load syllabus
        syllabus = load_syllabus(subject, grade)
        
        # Filter topics for term
        term_topics = [t for t in syllabus if str(t['term']) == str(term)]
        
        if not term_topics:
            raise ValueError(f"No topics found for {subject} {grade} term {term}")
            
        # Create AI prompt with more explicit instructions
        prompt = f"""
        Create a weekly scheme of work for {subject} {grade} Term {term} 
        from week {start_week} to {end_week}.
        
        Use these topics as reference:
        {json.dumps(term_topics, indent=2)}
        
        You must return a valid JSON array containing exactly {end_week - start_week + 1} weeks.
        Each object in the array must follow this exact structure:
        {{
            "week": <week number>,
            "periods": <number between 2-4>,
            "theme": "<main theme>",
            "topic": "<specific topic>",
            "sub_topic": "<specific sub-topic>",
            "competency": "<what students should be able to do>",
            "learning_outcomes": ["outcome 1", "outcome 2", ...],
            "activities": ["activity 1", "activity 2", ...],
            "methodology": ["method 1", "method 2", ...],
            "resources": ["resource 1", "resource 2", ...],
            "references": ["reference 1", "reference 2", ...],
            "remarks": "<additional notes>"
        }}

        Important:
        1. Return ONLY the JSON array, no other text
        2. Ensure all fields are present
        3. All string arrays must have at least one item
        4. Week numbers must be sequential from {start_week} to {end_week}
        5. All text fields must be non-empty strings
        """
        
        # Get AI response
        model = get_ai_model()
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Log the raw response for debugging
        logging.debug(f"Raw AI response: {response_text[:500]}...")
        
        # Parse and validate response
        try:
            # Try to extract JSON array if response contains extra text
            import re
            json_match = re.search(r'\[(.*)\]', response_text, re.DOTALL)
            if json_match:
                response_text = f"[{json_match.group(1)}]"
            
            topics = json.loads(response_text)
            
            # Validate response structure
            if not isinstance(topics, list):
                raise ValueError("Response must be a JSON array")
                
            if len(topics) != (end_week - start_week + 1):
                raise ValueError(f"Expected {end_week - start_week + 1} weeks, got {len(topics)}")
                
            # Validate each topic
            required_fields = ['week', 'periods', 'theme', 'topic', 'sub_topic', 'competency', 
                             'learning_outcomes', 'activities', 'methodology', 'resources', 
                             'references', 'remarks']
            
            for topic in topics:
                if not isinstance(topic, dict):
                    raise ValueError("Each item must be a dictionary")
                    
                # Check all required fields are present
                missing_fields = [field for field in required_fields if field not in topic]
                if missing_fields:
                    raise ValueError(f"Missing required fields: {missing_fields}")
                    
                # Validate field types
                if not isinstance(topic['week'], (int, str)) or not str(topic['week']).isdigit():
                    raise ValueError("Week must be a number")
                if not isinstance(topic['periods'], (int, str)) or not str(topic['periods']).isdigit():
                    raise ValueError("Periods must be a number")
                    
                # Validate arrays have content
                array_fields = ['learning_outcomes', 'activities', 'methodology', 'resources', 'references']
                for field in array_fields:
                    if not isinstance(topic[field], list) or len(topic[field]) == 0:
                        raise ValueError(f"{field} must be a non-empty array")
            
            return topics
            
        except json.JSONDecodeError as e:
            logging.error(f"JSON decode error: {str(e)}")
            logging.error(f"Response text: {response_text}")
            raise ValueError("Invalid JSON format in AI response")
            
    except Exception as e:
        logging.error(f"Error generating scheme topics: {str(e)}")
        raise ValueError(f"Failed to generate scheme: {str(e)}")

def load_syllabus(subject, grade):
    """Load syllabus data from file"""
    try:
        path = os.path.join('static', 'syllabi', subject, f"{grade}.json")
        with open(path) as f:
            return json.load(f)
    except Exception as e:
        raise ValueError(f"Could not load syllabus for {subject} {grade}: {e}")

def generate_teaching_content(subject, topic, grade_level, options):
    """Generate teaching content using AI based on type and options"""
    try:
        content_type = options.get('content_type')
        
        # Build the base prompt based on content type
        if content_type == 'lesson':
            prompt = _generate_lesson_prompt(subject, topic, grade_level, options)
        elif content_type == 'visual':
            prompt = _generate_visual_prompt(subject, topic, grade_level, options)
        elif content_type == 'activity':
            prompt = _generate_activity_prompt(subject, topic, grade_level, options)
        else:
            raise ValueError(f"Invalid content type: {content_type}")
            
        # Get AI response
        model = get_ai_model()
        response = model.generate_content(prompt)
        
        # Parse and validate response
        try:
            content = json.loads(response.text)
        except json.JSONDecodeError:
            # Try to extract JSON from markdown code block
            import re
            json_match = re.search(r'```json\s*(.*?)\s*```', response.text, re.DOTALL)
            if json_match:
                content = json.loads(json_match.group(1))
            else:
                raise ValueError("Invalid JSON response from AI")
                
        return content
        
    except Exception as e:
        current_app.logger.error(f"Error in generate_teaching_content: {str(e)}")
        current_app.logger.error(f"Options: {options}")
        raise

def _generate_lesson_prompt(subject, topic, grade_level, options):
    """Generate prompt for lesson content"""
    duration = options.get('duration', 40)
    components = [k.replace('include_', '') for k in options if k.startswith('include_')]
    
    return f"""
    Create a detailed lesson plan for teaching {topic} in {subject} at {grade_level} level.
    
    Duration: {duration} minutes
    
    Include the following components: {', '.join(components)}
    
    Return the response in this JSON structure:
    {{
        "title": "Lesson title",
        "duration": "{duration} minutes",
        "components": {{
            "introduction": {{
                "prior_knowledge": ["list of required prior knowledge"],
                "hook": "engaging introduction activity",
                "duration": "time in minutes"
            }},
            "objectives": {{
                "learning_outcomes": ["list of specific learning outcomes"],
                "success_criteria": ["list of success criteria"]
            }},
            "main_content": {{
                "key_concepts": ["list of key concepts"],
                "explanations": ["detailed explanations"],
                "examples": ["relevant examples"]
            }},
            "activities": [
                {{
                    "title": "activity name",
                    "type": "individual/group",
                    "duration": "time in minutes",
                    "description": "activity details",
                    "materials": ["required materials"],
                    "steps": ["step-by-step instructions"]
                }}
            ],
            "assessment": {{
                "questions": ["assessment questions"],
                "answers": ["model answers"],
                "rubric": ["assessment criteria"]
            }}
        }},
        "teaching_notes": ["important teaching notes"],
        "resources": ["list of additional resources"],
        "differentiation": {{
            "support": ["support strategies"],
            "extension": ["extension activities"]
        }}
    }}
    """

def _generate_visual_prompt(subject, topic, grade_level, options):
    """Generate prompt for visual content"""
    visual_type = options.get('visual_type')
    complexity = options.get('complexity')
    include_labels = options.get('include_labels', True)
    
    return f"""
    Create a detailed {visual_type} for teaching {topic} in {subject} at {grade_level} level.
    
    Complexity level: {complexity}
    Include labels and annotations: {include_labels}
    
    Return the response in this JSON structure:
    {{
        "title": "Visual title",
        "type": "{visual_type}",
        "description": "detailed description of the visual",
        "components": [
            {{
                "element": "element name",
                "description": "what this element represents",
                "labels": ["list of labels"],
                "connections": ["relationships to other elements"]
            }}
        ],
        "teaching_notes": {{
            "setup": "how to present this visual",
            "key_points": ["important points to emphasize"],
            "common_misconceptions": ["potential misunderstandings to address"]
        }},
        "student_activities": [
            {{
                "type": "activity type",
                "description": "how students should interact with the visual",
                "questions": ["guiding questions"]
            }}
        ]
    }}
    """

def _generate_activity_prompt(subject, topic, grade_level, options):
    """Generate prompt for activity content"""
    activity_type = options.get('activity_type')
    duration = options.get('activity_duration', 20)
    components = [k.replace('include_', '') for k in options if k.startswith('include_')]
    
    return f"""
    Create a detailed {activity_type} activity for teaching {topic} in {subject} at {grade_level} level.
    
    Duration: {duration} minutes
    Include these components: {', '.join(components)}
    
    Return the response in this JSON structure:
    {{
        "title": "Activity title",
        "type": "{activity_type}",
        "duration": "{duration} minutes",
        "objectives": ["learning objectives"],
        "materials": ["required materials"],
        "setup": {{
            "preparation": ["preparation steps"],
            "classroom_arrangement": "required arrangement",
            "safety_considerations": ["safety notes"]
        }},
        "procedure": {{
            "introduction": "how to introduce the activity",
            "steps": ["detailed step-by-step instructions"],
            "timing": ["time allocation for each step"]
        }},
        "student_instructions": {{
            "task_description": "clear task description",
            "steps": ["student steps to follow"],
            "deliverables": ["what students should produce"]
        }},
        "assessment": {{
            "criteria": ["assessment criteria"],
            "rubric": [
                {{
                    "criterion": "what to assess",
                    "levels": ["level descriptions"],
                    "marks": ["mark allocation"]
                }}
            ]
        }},
        "extensions": ["extension activities"],
        "troubleshooting": ["common issues and solutions"]
    }}
    """

def generate_question_variations(question_data, num_variations=2):
    """Generate variations of a question using AI."""
    try:
        model = get_ai_model()
        
        # Define the JSON structure separately to avoid nested f-strings
        json_structure = r'''{
            "variations": [
                {
                    "question_text": "variation text here",
                    "context": "background or scenario",
                    "solution": "correct answer or solution steps",
                    "options": [
                        {"text": "option text", "is_correct": true/false}
                    ],
                    "explanation": "why this tests the same concept"
                }
            ]
        }'''
        
        prompt = f"""
        Generate {num_variations} variations of this {question_data['subject']} question:
        
        Original Question: {question_data['content']}
        Topic: {question_data['topic']}
        Type: {question_data['question_type']}
        
        Requirements:
        1. Maintain the same difficulty level and learning objective
        2. Use different scenarios or contexts
        3. For numerical questions:
           - Use different values but similar relationships
           - Maintain similar complexity in calculations
        4. For conceptual questions:
           - Apply the same concept to different situations
           - Use varied real-world examples
        
        Return the response in this JSON structure:
        {json_structure}
        """
        
        response = model.generate_content(prompt)
        
        # Parse and validate JSON response
        try:
            variations_data = json.loads(response.text)
        except json.JSONDecodeError:
            import re
            json_match = re.search(r'```json\s*(.*?)\s*```', response.text, re.DOTALL)
            if json_match:
                variations_data = json.loads(json_match.group(1))
            else:
                raise ValueError("No valid JSON found in response")
        
        return variations_data['variations']
        
    except Exception as e:
        current_app.logger.error(f"Error generating question variations: {str(e)}")
        return []


